import argparse
import copy
import json
import re
import shutil
import subprocess
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import List, Tuple, Dict, Any

from num2words import num2words

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PLACEHOLDERS = {
    "buyer": ["{{buyer}}", "{{company_name}}", "{{customer}}"],
    "service": ["{{service}}", "{{invoice_for}}", "{{item}}"],
    "amount": ["{{amount}}", "{{sum}}"],
    "amount_words": ["{{amount_words}}", "{{total_sum_words}}"],
    "contract_basis": ["{{contract_basis}}", "{{basis}}"],
    "invoice_number": ["{{invoice_number}}", "{{invoice_no}}"],
    "invoice_date": ["{{invoice_date}}", "{{date}}"],
    "tax_amount": ["{{tax_amount}}", "{{tax_5}}"],
    "total_with_tax": ["{{total_with_tax}}", "{{amount_with_tax}}"],
    "total_with_tax_words": [
        "{{total_with_tax_words}}",
        "{{amount_with_tax_words}}",
    ],
}

BOLD_PLACEHOLDERS = {
    "buyer": ["{{buyer_b}}", "{{company_name_b}}", "{{customer_b}}"],
    "service": ["{{service_b}}", "{{invoice_for_b}}", "{{item_b}}"],
    "amount": ["{{amount_b}}", "{{sum_b}}"],
    "amount_words": ["{{amount_words_b}}", "{{total_sum_words_b}}"],
    "contract_basis": ["{{contract_basis_b}}", "{{basis_b}}"],
    "invoice_number": ["{{invoice_number_b}}", "{{invoice_no_b}}"],
    "invoice_date": ["{{invoice_date_b}}", "{{date_b}}"],
    "tax_amount": ["{{tax_amount_b}}", "{{tax_5_b}}"],
    "total_with_tax": ["{{total_with_tax_b}}", "{{amount_with_tax_b}}"],
    "total_with_tax_words": [
        "{{total_with_tax_words_b}}",
        "{{amount_with_tax_words_b}}",
    ],
}

ITEM_PLACEHOLDERS = [
    "{{item_no}}",
    "{{item_name}}",
    "{{item_qty}}",
    "{{item_unit}}",
    "{{item_price}}",
    "{{item_sum}}",
]


def parse_decimal(value: str) -> Decimal:
    cleaned = value.strip().replace("\u00a0", " ").replace(" ", "")
    cleaned = cleaned.replace("₽", "").replace("руб.", "").replace("руб", "")
    if "," in cleaned and "." in cleaned:
        cleaned = cleaned.replace(".", "")
    cleaned = cleaned.replace(",", ".")
    cleaned = re.sub(r"[^\d.]", "", cleaned) or "0"
    return Decimal(cleaned)


def format_money(value: Decimal) -> str:
    quantized = value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    parts = f"{quantized:.2f}".split(".")
    rubles = parts[0]
    kopeks = parts[1]
    rubles_spaced = " ".join(
        [rubles[max(i - 3, 0):i] for i in range(len(rubles), 0, -3)][::-1]
    )
    return f"{rubles_spaced},{kopeks}"


def apply_run_style(new_run, base_run, force_bold: bool | None = None) -> None:
    if base_run.style is not None:
        new_run.style = base_run.style
    new_run.font.name = base_run.font.name
    new_run.font.size = base_run.font.size
    new_run.font.color.rgb = base_run.font.color.rgb
    new_run.font.italic = base_run.font.italic
    new_run.font.underline = base_run.font.underline
    if force_bold is None:
        new_run.font.bold = base_run.font.bold
    else:
        new_run.font.bold = force_bold


def replace_in_paragraph(paragraph, replacements: dict, bold_replacements: dict) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace(key, value)
    for key, value in bold_replacements.items():
        new_text = new_text.replace(key, f"[[B]]{value}[[/B]]")

    if new_text == full_text:
        return

    base_run = paragraph.runs[0]
    for run in paragraph.runs:
        run.text = ""

    parts = new_text.split("[[/B]]")
    for part in parts:
        if "[[B]]" in part:
            before, bold_text = part.split("[[B]]", 1)
            if before:
                run = paragraph.add_run(before)
                apply_run_style(run, base_run)
            bold_run = paragraph.add_run(bold_text)
            apply_run_style(bold_run, base_run, force_bold=True)
        else:
            if part:
                run = paragraph.add_run(part)
                apply_run_style(run, base_run)


def replace_in_table(table, replacements: dict, bold_replacements: dict) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements, bold_replacements)


def build_replacements(
    buyer: str,
    service: str,
    amount: str,
    contract_basis: str,
    invoice_number: str,
    invoice_date: str,
    tax_amount: str = "",
    total_with_tax: str = "",
    total_with_tax_words: str = "",
) -> dict:
    mapping = {}
    for key in PLACEHOLDERS["buyer"]:
        mapping[key] = buyer
    for key in PLACEHOLDERS["service"]:
        mapping[key] = service
    for key in PLACEHOLDERS["amount"]:
        mapping[key] = amount
    amount_words = amount_to_words(amount)
    for key in PLACEHOLDERS["amount_words"]:
        mapping[key] = amount_words
    for key in PLACEHOLDERS["contract_basis"]:
        mapping[key] = contract_basis
    for key in PLACEHOLDERS["invoice_number"]:
        mapping[key] = invoice_number
    for key in PLACEHOLDERS["invoice_date"]:
        mapping[key] = invoice_date
    for key in PLACEHOLDERS["tax_amount"]:
        mapping[key] = tax_amount
    for key in PLACEHOLDERS["total_with_tax"]:
        mapping[key] = total_with_tax
    for key in PLACEHOLDERS["total_with_tax_words"]:
        mapping[key] = total_with_tax_words
    return mapping


def build_bold_replacements(
    buyer: str,
    service: str,
    amount: str,
    contract_basis: str,
    invoice_number: str,
    invoice_date: str,
    tax_amount: str = "",
    total_with_tax: str = "",
    total_with_tax_words: str = "",
) -> dict:
    mapping = {}
    for key in BOLD_PLACEHOLDERS["buyer"]:
        mapping[key] = buyer
    for key in BOLD_PLACEHOLDERS["service"]:
        mapping[key] = service
    for key in BOLD_PLACEHOLDERS["amount"]:
        mapping[key] = amount
    amount_words = amount_to_words(amount)
    for key in BOLD_PLACEHOLDERS["amount_words"]:
        mapping[key] = amount_words
    for key in BOLD_PLACEHOLDERS["contract_basis"]:
        mapping[key] = contract_basis
    for key in BOLD_PLACEHOLDERS["invoice_number"]:
        mapping[key] = invoice_number
    for key in BOLD_PLACEHOLDERS["invoice_date"]:
        mapping[key] = invoice_date
    for key in BOLD_PLACEHOLDERS["tax_amount"]:
        mapping[key] = tax_amount
    for key in BOLD_PLACEHOLDERS["total_with_tax"]:
        mapping[key] = total_with_tax
    for key in BOLD_PLACEHOLDERS["total_with_tax_words"]:
        mapping[key] = total_with_tax_words
    return mapping


def amount_to_words(amount: str) -> str:
    # Accept formats like "50 000,00", "50000.00", "50000"
    value = parse_decimal(amount).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    rubles = int(value)
    kopeks = int((value - Decimal(rubles)) * 100)

    rub_words = num2words(rubles, lang="ru")
    # Capitalize first letter, standard in documents
    rub_words = rub_words[:1].upper() + rub_words[1:]
    return f"{rub_words} рублей {kopeks:02d} копеек"


def format_invoice_date(raw_date: str) -> str:
    if not raw_date:
        raw_date = datetime.now().strftime("%d.%m.%Y")
    raw_date = raw_date.strip()
    for fmt in ("%d.%m.%Y", "%Y-%m-%d"):
        try:
            dt = datetime.strptime(raw_date, fmt)
            break
        except ValueError:
            dt = None
    if dt is None:
        # Assume already formatted like "02 февраля 2026"
        return raw_date
    months = [
        "января",
        "февраля",
        "марта",
        "апреля",
        "мая",
        "июня",
        "июля",
        "августа",
        "сентября",
        "октября",
        "ноября",
        "декабря",
    ]
    return f"{dt.day:02d} {months[dt.month - 1]} {dt.year}"


def load_counter(counter_path: Path) -> dict:
    if not counter_path.exists():
        return {"date": "", "numbers": {"ooo": 0, "ip": 0}}
    try:
        data = json.loads(counter_path.read_text(encoding="utf-8"))
        # Backward compatibility with old flat format
        if "numbers" not in data:
            number = int(data.get("number", 0))
            return {"date": data.get("date", ""), "numbers": {"ooo": number, "ip": 0}}
        if "ooo" not in data["numbers"]:
            data["numbers"]["ooo"] = 0
        if "ip" not in data["numbers"]:
            data["numbers"]["ip"] = 0
        return data
    except json.JSONDecodeError:
        return {"date": "", "numbers": {"ooo": 0, "ip": 0}}


def next_invoice_number(counter_path: Path, invoice_type: str) -> str:
    today = datetime.now().strftime("%Y-%m-%d")
    state = load_counter(counter_path)
    if state.get("date") != today:
        state["date"] = today
        state["numbers"] = {"ooo": 0, "ip": 0}
    invoice_type_key = "ip" if invoice_type.lower() == "ip" else "ooo"
    state["numbers"][invoice_type_key] = int(state["numbers"].get(invoice_type_key, 0)) + 1
    counter_path.write_text(json.dumps(state, ensure_ascii=False), encoding="utf-8")
    return str(state["numbers"][invoice_type_key])


def parse_items(raw_items: str) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    if not raw_items:
        return items
    for idx, raw in enumerate(raw_items.split(";"), start=1):
        if not raw.strip():
            continue
        parts = [p.strip() for p in raw.split("|")]
        if len(parts) != 4:
            raise ValueError(
                "Каждый товар должен иметь 4 поля: "
                "название|кол-во|ед.|цена"
            )
        name, qty, unit, price = parts
        qty_value = parse_decimal(qty)
        price_value = parse_decimal(price)
        item_sum = (qty_value * price_value).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
        items.append(
            {
                "name": name,
                "qty_raw": qty,
                "unit": unit,
                "price_value": price_value,
                "sum_value": item_sum,
                "price_fmt": format_money(price_value),
                "sum_fmt": format_money(item_sum),
            }
        )
    return items


def row_contains_item_placeholders(row) -> bool:
    row_text = " ".join(
        p.text for cell in row.cells for p in cell.paragraphs if p.text
    )
    return any(ph in row_text for ph in ITEM_PLACEHOLDERS)


def fill_item_row(row, item_no: str, item: Dict[str, Any]) -> None:
    replacements = {
        "{{item_no}}": item_no,
        "{{item_name}}": item["name"],
        "{{item_qty}}": item["qty_raw"],
        "{{item_unit}}": item["unit"],
        "{{item_price}}": item["price_fmt"],
        "{{item_sum}}": item["sum_fmt"],
    }
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            replace_in_paragraph(paragraph, replacements, {})


def row_from_tr(table, tr):
    for row in table.rows:
        if row._tr is tr:
            return row
    return None


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, spec in kwargs.items():
        edge_tag = qn(f"w:{edge}")
        element = borders.find(edge_tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for key, value in spec.items():
            element.set(qn(f"w:{key}"), str(value))


def normalize_table_inner_borders(table) -> None:
    # Make inner horizontal borders thin while keeping outer borders.
    row_count = len(table.rows)
    if row_count <= 1:
        return
    for r_idx, row in enumerate(table.rows):
        is_last = r_idx == row_count - 1
        if is_last:
            continue
        for cell in row.cells:
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": "4", "space": "0", "color": "000000"},
            )


def expand_items_in_table(table, items: List[Dict[str, Any]]) -> bool:
    if not items:
        return False
    template_row = None
    template_index = None
    for idx, row in enumerate(table.rows):
        if row_contains_item_placeholders(row):
            template_row = row
            template_index = idx
            break
    if template_row is None or template_index is None:
        return False

    template_tr = template_row._tr
    insert_after = template_tr

    for i, item in enumerate(items, start=1):
        new_tr = copy.deepcopy(template_tr)
        insert_after.addnext(new_tr)
        new_row = row_from_tr(table, new_tr)
        if new_row is None:
            continue
        fill_item_row(new_row, str(i), item)
        insert_after = new_tr

    template_tr.getparent().remove(template_tr)
    return True


def generate_invoice_pdf(
    *,
    template_path: Path,
    output_docx: Path,
    output_pdf: Path,
    buyer: str,
    contract_basis: str,
    item_name: str,
    item_price: str,
    item_qty: str = "1",
    item_unit: str = "усл.",
    service: str = "",
    amount: str = "",
    items: str = "",
    invoice_type: str = "ooo",
    invoice_number: str = "auto",
    invoice_date: str = "",
    counter_file: Path = Path(".invoice_counter.json"),
) -> None:
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template_path))
    if items:
        items_list = parse_items(items)
    elif item_name and item_price:
        single = f"{item_name}|{item_qty}|{item_unit}|{item_price}"
        items_list = parse_items(single)
    else:
        items_list = []

    if items_list:
        total = sum((item["sum_value"] for item in items_list), Decimal("0.00"))
        amount_value = format_money(total)
    else:
        if not amount:
            raise ValueError("Нужно указать amount или item/items.")
        amount_value = amount

    formatted_date = format_invoice_date(invoice_date)
    if invoice_number == "auto":
        invoice_number_value = next_invoice_number(counter_file, invoice_type)
    else:
        invoice_number_value = invoice_number

    service_value = service or (items_list[0]["name"] if items_list else "")

    tax_amount_value = ""
    total_with_tax_value = ""
    total_with_tax_words_value = ""
    if invoice_type.lower() == "ip":
        amount_decimal = parse_decimal(amount_value)
        tax_amount = (amount_decimal * Decimal("0.05")).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
        total_with_tax = (amount_decimal + tax_amount).quantize(
            Decimal("0.01"), rounding=ROUND_HALF_UP
        )
        tax_amount_value = format_money(tax_amount)
        total_with_tax_value = format_money(total_with_tax)
        total_with_tax_words_value = amount_to_words(total_with_tax_value)

    replacements = build_replacements(
        buyer,
        service_value,
        amount_value,
        contract_basis,
        invoice_number_value,
        formatted_date,
        tax_amount_value,
        total_with_tax_value,
        total_with_tax_words_value,
    )
    bold_replacements = build_bold_replacements(
        buyer,
        service_value,
        amount_value,
        contract_basis,
        invoice_number_value,
        formatted_date,
        tax_amount_value,
        total_with_tax_value,
        total_with_tax_words_value,
    )

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements, bold_replacements)
    for table in doc.tables:
        replace_in_table(table, replacements, bold_replacements)
        changed = expand_items_in_table(table, items_list)
        if changed:
            normalize_table_inner_borders(table)

    if output_docx.exists():
        output_docx.unlink()
    doc.save(str(output_docx))

    if output_pdf.exists():
        output_pdf.unlink()
    convert_docx_to_pdf(output_docx, output_pdf)


def convert_docx_to_pdf(input_docx: Path, output_pdf: Path) -> None:
    soffice_path = shutil.which("soffice")
    if not soffice_path:
        raise RuntimeError(
            "Не найден LibreOffice (soffice). Установи libreoffice и убедись, "
            "что команда soffice доступна в PATH."
        )

    outdir = output_pdf.parent
    outdir.mkdir(parents=True, exist_ok=True)

    command = [
        soffice_path,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(outdir),
        str(input_docx),
    ]

    result = subprocess.run(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        check=False,
    )

    if result.returncode != 0:
        raise RuntimeError(
            "Ошибка конвертации DOCX в PDF через LibreOffice.\n"
            f"stdout: {result.stdout}\n"
            f"stderr: {result.stderr}"
        )

    generated_pdf = outdir / f"{input_docx.stem}.pdf"
    if not generated_pdf.exists():
        raise RuntimeError(
            "LibreOffice завершился без ошибки, но PDF не найден. "
            "Проверь вывод команды soffice."
        )

    if generated_pdf.resolve() != output_pdf.resolve():
        if output_pdf.exists():
            output_pdf.unlink()
        generated_pdf.replace(output_pdf)


def main() -> None:
    parser = argparse.ArgumentParser(description="Fill DOCX invoice and export to PDF.")
    parser.add_argument("--template", default="invoice.docx", help="Путь к DOCX шаблону")
    parser.add_argument("--output-docx", default="invoice_filled.docx", help="Выходной DOCX")
    parser.add_argument("--output-pdf", default="invoice.pdf", help="Выходной PDF")
    parser.add_argument("--buyer", required=True, help="Название покупателя")
    parser.add_argument(
        "--invoice-type",
        default="ooo",
        choices=["ooo", "ip"],
        help="Тип счета: ooo или ip",
    )
    parser.add_argument("--service", default="", help="Название услуги (если не используешь item)")
    parser.add_argument("--amount", default="", help="Сумма (если не используешь items)")
    parser.add_argument("--contract-basis", default="", help="Основание договора")
    parser.add_argument(
        "--invoice-number",
        default="auto",
        help='Номер счета ("auto" для автонумерации)',
    )
    parser.add_argument(
        "--invoice-date",
        default="",
        help='Дата счета (например "02.02.2026" или "2026-02-02")',
    )
    parser.add_argument("--item-name", default="", help="Название товара (если один)")
    parser.add_argument("--item-qty", default="1", help="Кол-во товара (если один)")
    parser.add_argument("--item-unit", default="усл.", help="Ед. измерения (если один)")
    parser.add_argument("--item-price", default="", help="Цена товара (если один)")
    parser.add_argument(
        "--items",
        default="",
        help="Список товаров: 'название|кол-во|ед.|цена;...'",
    )
    parser.add_argument(
        "--counter-file",
        default=".invoice_counter.json",
        help="Файл для хранения автонумерации",
    )

    args = parser.parse_args()
    generate_invoice_pdf(
        template_path=Path(args.template),
        output_docx=Path(args.output_docx),
        output_pdf=Path(args.output_pdf),
        buyer=args.buyer,
        contract_basis=args.contract_basis,
        item_name=args.item_name,
        item_price=args.item_price,
        item_qty=args.item_qty,
        item_unit=args.item_unit,
        service=args.service,
        amount=args.amount,
        items=args.items,
        invoice_type=args.invoice_type,
        invoice_number=args.invoice_number,
        invoice_date=args.invoice_date,
        counter_file=Path(args.counter_file),
    )


if __name__ == "__main__":
    main()
